from django.views.generic.edit import CreateView
from .models import Booking
from docx import Document

class BookingCreate(CreateView):
    model = Booking
    fields = ['name']
    success_url = '/'
    def create_doc(self):
        doc = Document()
        doc.add_heading('This is the title', 0)
        p = doc.add_paragraph(Booking.name)
        doc.save('serenitysoft/output')
